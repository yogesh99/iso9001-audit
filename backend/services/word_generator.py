from docx import Document
from backend.services.summary_engine import generate_report_summary
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

import os

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")

os.makedirs(OUTPUT_DIR, exist_ok=True)


def clear_cell(cell):
    """Remove existing paragraphs safely"""
    cell._tc.clear_content()
    

SECTION_HEADINGS = [
    "Evidences and comments on conformance:",
    "References to documentation:",
    "Findings/NCs:"
]

# Heading to remove: its content is merged under "Evidences and comments on conformance:"
AUDITOR_EVIDENCE_HEADING = "Auditor Evidence (site specific observations):"

# Default/system narrative lines to never include in the report (auditor-only content)
DEFAULT_NARRATIVE_LINES = {
    "The organization has identified internal issues relevant to its purpose and strategic direction.",
    "External issues relevant to the organization have been identified.",
    "A documented mechanism is established to monitor and review these issues.",
    "The organization has identified internal issues relevant to its purpose and strategic direction. External issues relevant to the organization have been identified. A documented mechanism is established to monitor and review these issues.",
}

def add_heading_paragraph(cell, text, template_para):
    p = cell.add_paragraph()
    run = p.add_run(text)

    if template_para.runs:
        run.font.name = template_para.runs[0].font.name
        run.font.size = template_para.runs[0].font.size

    run.bold = True

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1

def add_paragraph(cell, text, template_para):
    p = cell.add_paragraph()
    run = p.add_run(text)

    # Font
    if template_para.runs:
        run.font.name = template_para.runs[0].font.name
        run.font.size = template_para.runs[0].font.size

    run.bold = False
    run.italic = False

    # 🔴 SPACING FIX
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1

    
    

def add_small_paragraph(cell, text, template_para, font_size_pt=9):
    p = cell.add_paragraph()
    run = p.add_run(text)

    if template_para.runs:
        run.font.name = template_para.runs[0].font.name

    run.font.size = Pt(font_size_pt)
    run.bold = False
    run.italic = False

    p.paragraph_format.space_before = template_para.paragraph_format.space_before
    p.paragraph_format.space_after = template_para.paragraph_format.space_after
    p.paragraph_format.line_spacing = template_para.paragraph_format.line_spacing


def insert_paragraph_before(paragraph):
    """
    Insert a paragraph before the given paragraph (version-safe).
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_standard_at_start(doc, standard_name: str):
    """
    Insert a visible 'Audit Standard: ISO 9001' (or ISO 13485) line at the very
    start of the document so it shows in every downloaded report.
    """
    body = doc.element.body
    if len(body) == 0:
        return
    first = body[0]
    new_p = OxmlElement("w:p")
    first.addprevious(new_p)
    para = Paragraph(new_p, body)
    run = para.add_run(f"Audit Standard: {standard_name}")
    run.bold = True
    run.font.size = Pt(12)



def fill_clause_in_tables(doc, clause_id, clause_data):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells

            if cells[0].text.strip() == clause_id:

                evidence_cell = cells[1]
                status_cell = cells[2]

                template_para = evidence_cell.paragraphs[0]

                clear_cell(evidence_cell)
                clear_cell(status_cell)

                # ---- Evidences and comments on conformance ----
                add_heading_paragraph(
                    evidence_cell,
                    "Evidences and comments on conformance:",
                    template_para
                )

                for line in clause_data["evidence"].split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    # Omit this title; its content stays under "Evidences and comments on conformance:"
                    if line == AUDITOR_EVIDENCE_HEADING:
                        continue
                    # Never output default/system narrative in the report
                    if line in DEFAULT_NARRATIVE_LINES or any(
                        default in line for default in DEFAULT_NARRATIVE_LINES
                    ):
                        continue

                    if line in SECTION_HEADINGS:
                        add_heading_paragraph(evidence_cell, line, template_para)
                    else:
                        add_paragraph(evidence_cell, line, template_para)

                # ---- References ----
                add_heading_paragraph(
                    evidence_cell,
                    "References to documentation:",
                    template_para
                )
                add_paragraph(evidence_cell, clause_data["documents"], template_para)

                # ---- Findings / NCs ----
                add_heading_paragraph(
                    evidence_cell,
                    "Findings/NCs:",
                    template_para
                )
                add_paragraph(evidence_cell, clause_data["findings"], template_para)

                # ---- Status ----
                p = status_cell.add_paragraph()
                run = p.add_run(clause_data["status"])

                if template_para.runs:
                    run.font.name = template_para.runs[0].font.name
                    run.font.size = template_para.runs[0].font.size

                return





def generate_report(audit_data):
    audit_type = audit_data.get("audit_type", "Stage 2")
    if audit_type == "Stage 1":
        template_name = "stage1_template.docx"
    elif audit_type == "Surveillance":
        template_name = "surveillance_template.docx"
    else:
        template_name = "stage2_template.docx"
    # Fallback if surveillance template is missing (use Stage 2)
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    if not os.path.isfile(template_path) and audit_type == "Surveillance":
        template_path = os.path.join(TEMPLATE_DIR, "stage2_template.docx")
    doc = Document(template_path)

    standard = audit_data.get("standard", "iso9001")
    standard_name = "ISO 13485" if standard == "iso13485" else "ISO 9001"
    insert_standard_at_start(doc, standard_name)

    for clause_id, clause_data in audit_data["clauses"].items():
        fill_clause_in_tables(doc, clause_id, clause_data)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    client_name = audit_data.get("client_name", "Client")
    audit_type = audit_data.get("audit_type", "Audit")
    safe_client = client_name.replace(" ", "_")
    safe_audit = audit_type.replace(" ", "_")
    safe_standard = "ISO13485" if standard == "iso13485" else "ISO9001"

    output_path = os.path.join(
        OUTPUT_DIR,
        f"{safe_client}_{safe_audit}_{safe_standard}_{timestamp}.docx"
    )

    summary = generate_report_summary(audit_data["clauses"])
    fill_report_summary(doc, summary, standard_name=standard_name)

    doc.save(output_path)
    return output_path


from docx.shared import Pt

def fill_report_summary(doc, summary: dict, standard_name: str = "ISO 9001"):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:

                if "Report Summary" in cell.text:

                    template_para = cell.paragraphs[0]

                    # Find the NOTE paragraph (must exist in template)
                    note_para = None
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if p.text.strip().startswith("Note:"):
                            note_para = p
                            break

                    if not note_para:
                        return  # safety exit

                    def add_summary_line(text, bold=False):
                        p = insert_paragraph_before(note_para)
                        run = p.add_run(text)

                        if template_para.runs:
                            run.font.name = template_para.runs[0].font.name

                        run.bold = bold
                        run.font.size = Pt(9)

                    # ---- Insert content ----
                    add_summary_line(f"Standard: {standard_name}", bold=True)
                    add_summary_line("No nonconformities identified." if not summary["nc_clauses"]
                                     else "Clauses with nonconformities: " + ", ".join(summary["nc_clauses"]))

                    add_summary_line(f"Not Applicable: {summary['Not Applicable']}")
                    add_summary_line(f"Major NCs: {summary['Major NC']}")
                    add_summary_line(f"Minor NCs: {summary['Minor NC']}")
                    add_summary_line(f"Compliant: {summary['Compliant']}")
                    add_summary_line(f"Total clauses audited: {summary['total']}")

                

                    return




